# ui/acronymswindow.py
from PyQt5.QtWidgets import (QMainWindow, QVBoxLayout, QHBoxLayout, QWidget,
                             QPushButton, QFileDialog, QSizePolicy,
                             QTableWidget, QTableWidgetItem, QCheckBox,
                             QTabWidget, QHeaderView, QMessageBox, QAbstractScrollArea)
import os
import sys
import requests
import win32com.client
from PyQt5.QtGui import QIcon
from docx import Document
from PyQt5.QtCore import Qt, QTimer

try:
    from macros.Acronyms import find_acronyms, get_definition
except ImportError:
    print("Error: Could not import Acronyms macro. Make sure macros/Acronyms.py exists.")
    find_acronyms = None
    get_definition = None

def fetch_acronym_list_online(url, base_cache_path):
    """
    Fetch the acronym list from the online URL.
    On success, write the file to base_cache_path and return the path.
    On failure, if a cached base file exists, return that;
    otherwise, raise an exception.
    """
    try:
        response = requests.get(url, timeout=10)
        response.raise_for_status()
        os.makedirs(os.path.dirname(base_cache_path), exist_ok=True)
        with open(base_cache_path, "w", encoding="utf-8") as f:
            f.write(response.text)
        print(f"Fetched base acronym list to: {base_cache_path}")
        return base_cache_path
    except Exception as e:
        print(f"Failed to fetch base acronym list online: {e}")
        if os.path.exists(base_cache_path):
            print(f"Using cached base acronym list: {base_cache_path}")
            return base_cache_path
        else:
            raise Exception(f"Failed to fetch base acronym list and no cache available: {e}")

class AcronymsWindow(QMainWindow):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.base_acronym_file_path = None
        self.user_acronym_file_path = os.path.join(os.path.expanduser("~"), ".doc_companion", "user_acronyms.txt")
        os.makedirs(os.path.dirname(self.user_acronym_file_path), exist_ok=True)

        self.setWindowTitle("Acronym Finder")
        self.setMinimumSize(800, 600)

        # --- Icon Paths ---
        dir_path = os.path.dirname(os.path.realpath(__file__))
        leaf_icon_path = os.path.join(dir_path, "leaf.ico")
        add_icon_path = os.path.join(dir_path, "add.ico")
        self.setWindowIcon(QIcon(leaf_icon_path))
        add_icon = QIcon(add_icon_path)
        # --- End Icon Paths --- 

        self.layout = QVBoxLayout()
        self.layout.setContentsMargins(10, 10, 10, 10)
        self.layout.setSpacing(10)

        self.run_macro_button = QPushButton("Find Acronyms", self)
        self.run_macro_button.clicked.connect(self.run_macro)
        self.layout.addWidget(self.run_macro_button)

        self.tab_widget = QTabWidget(self)
        self.layout.addWidget(self.tab_widget)

        self.likely_table = self.create_table()
        self.possible_table = self.create_table()
        self.unlikely_table = self.create_table()

        # Add tabs WITHOUT the '+' button inside
        self.tab_widget.addTab(self.create_tab_widget(self.likely_table), "Likely")
        self.tab_widget.addTab(self.create_tab_widget(self.possible_table), "Possible")
        self.tab_widget.addTab(self.create_tab_widget(self.unlikely_table), "Unlikely")

        # --- Bottom Controls Layout ---
        bottom_layout = QHBoxLayout()
        bottom_layout.setSpacing(15)

        # Generate Button (Left)
        self.generate_table_button = QPushButton("Generate Table", self)
        self.generate_table_button.clicked.connect(self.generate_table)
        bottom_layout.addWidget(self.generate_table_button)

        bottom_layout.addStretch() # Push center/right

        # Check/Uncheck Buttons (Center)
        check_button_layout = QHBoxLayout()
        check_button_layout.setSpacing(10) # Spacing between check buttons

        self.likely_check_all_button = QPushButton('Likely', self)
        self.likely_check_all_button.setToolTip("Likely: Check/Uncheck All")
        self.likely_check_all_button.clicked.connect(lambda: self.check_uncheck_all(self.likely_table))
        self.likely_check_all_button.setObjectName("CheckUncheckButton")
        check_button_layout.addWidget(self.likely_check_all_button)

        self.possible_check_all_button = QPushButton('Possible', self)
        self.possible_check_all_button.setToolTip("Possible: Check/Uncheck All")
        self.possible_check_all_button.clicked.connect(lambda: self.check_uncheck_all(self.possible_table))
        self.possible_check_all_button.setObjectName("CheckUncheckButton")
        check_button_layout.addWidget(self.possible_check_all_button)

        self.unlikely_check_all_button = QPushButton('Unlikely', self)
        self.unlikely_check_all_button.setToolTip("Unlikely: Check/Uncheck All")
        self.unlikely_check_all_button.clicked.connect(lambda: self.check_uncheck_all(self.unlikely_table))
        self.unlikely_check_all_button.setObjectName("CheckUncheckButton")
        check_button_layout.addWidget(self.unlikely_check_all_button)

        bottom_layout.addLayout(check_button_layout)

        bottom_layout.addStretch() # Push right

        # Add Button (Right)
        self.add_row_button_global = QPushButton(add_icon, "")
        self.add_row_button_global.setObjectName("AddButton") # Set object name for QSS
        self.add_row_button_global.setToolTip("Add new row to current tab")
        self.add_row_button_global.clicked.connect(self.add_row_to_current_tab)
        bottom_layout.addWidget(self.add_row_button_global)

        self.layout.addLayout(bottom_layout) # Add bottom bar to main layout
        # --- End Bottom Controls ---

        self.widget = QWidget()
        self.widget.setLayout(self.layout)
        self.setCentralWidget(self.widget)
        self.update_stay_on_top()

    def create_tab_widget(self, table):
        """Creates a widget for a tab, containing only a table."""
        layout = QVBoxLayout()
        layout.setContentsMargins(0, 5, 0, 0)
        layout.setSpacing(5)
        layout.addWidget(table)
        widget = QWidget()
        widget.setLayout(layout)
        return widget

    def create_table(self):
        """Creates and configures a QTableWidget."""
        table = QTableWidget(0, 4)
        table.setHorizontalHeaderLabels(["Acronym", "Definition", "Include", "Context"])
        table.verticalHeader().setVisible(False)
        table.setAlternatingRowColors(True)
        table.setSizePolicy(QSizePolicy.Expanding, QSizePolicy.Expanding)
        table.setWordWrap(True) # Ensure Word Wrap is True

        table.verticalHeader().setMinimumSectionSize(25)
        table.verticalHeader().setDefaultSectionSize(35)

        table.setColumnWidth(0, 110)
        table.setColumnWidth(1, 350)
        table.setColumnWidth(2, 60)
        table.horizontalHeader().setSectionResizeMode(3, QHeaderView.Stretch)
        table.horizontalHeader().setSectionResizeMode(1, QHeaderView.Interactive)
        table.horizontalHeader().setSectionResizeMode(0, QHeaderView.Interactive)
        table.horizontalHeader().setSectionResizeMode(2, QHeaderView.Fixed)

        table.setSizeAdjustPolicy(QAbstractScrollArea.AdjustToContentsOnFirstShow)

        return table

    def run_macro(self):
        if not find_acronyms or not get_definition:
            QMessageBox.critical(self, "Error", "Acronym functions are not loaded.")
            return
        try:
            word_app = win32com.client.Dispatch('Word.Application')
            if not word_app.Documents.Count:
                QMessageBox.warning(self, "Warning", "No active Word document found.")
                return

            url = (
                "https://raw.githubusercontent.com/IIDelta/Doc_Companion/"
                "main/acronyms/acronym%20list.txt"
            )
            doc_companion_dir = os.path.join(os.path.expanduser("~"), ".doc_companion")
            base_list_filename = "base_acronym_list.txt"
            self.base_acronym_file_path = os.path.join(doc_companion_dir, base_list_filename)
            self.base_acronym_file_path = fetch_acronym_list_online(url, self.base_acronym_file_path)

            acronyms = find_acronyms(word_app, self.base_acronym_file_path, self.user_acronym_file_path)

            self.populate_table(self.likely_table, acronyms.get('likely', {}), True)
            self.populate_table(self.possible_table, acronyms.get('possible', {}), True)
            self.populate_table(self.unlikely_table, acronyms.get('unlikely', {}), False)

        except Exception as e:
            QMessageBox.critical(self, "Error", f"An error occurred: {str(e)}")
            print(f"Error running macro: {e}")

    def populate_table(self, table, data, is_checked_default):
        """Populates a table with acronym data and resizes rows."""
        table.setUpdatesEnabled(False) # Disable updates for speed
        table.setRowCount(0)
        for acronym, context in data.items():
            row_position = table.rowCount()
            table.insertRow(row_position)
            table.setItem(row_position, 0, QTableWidgetItem(acronym))
            definition_text = get_definition(acronym, self.base_acronym_file_path, self.user_acronym_file_path)
            table.setItem(row_position, 1, QTableWidgetItem(definition_text))

            checkbox = QCheckBox()
            checkbox.setChecked(is_checked_default)
            cell_widget = QWidget()
            cell_layout = QHBoxLayout(cell_widget)
            cell_layout.addWidget(checkbox)
            cell_layout.setAlignment(Qt.AlignCenter)
            cell_layout.setContentsMargins(0, 0, 0, 0)
            table.setCellWidget(row_position, 2, cell_widget)

            table.setItem(row_position, 3, QTableWidgetItem(context))

        table.setUpdatesEnabled(True) # Re-enable updates
        table.resizeRowsToContents() # Resize rows AFTER populating

    def generate_table(self):
        self.save_user_definitions()
        self.create_word_table()

    def save_user_definitions(self):
        user_definitions = {}
        try:
            with open(self.user_acronym_file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if '\t' in line:
                        acr, deph = line.split('\t', 1)
                        user_definitions[acr] = deph
        except FileNotFoundError:
            pass
        except Exception as e:
            QMessageBox.warning(self, "Warning", f"Could not read user definitions: {e}")
            return

        made_changes = False
        for table_widget in [self.likely_table, self.possible_table, self.unlikely_table]:
            for r in range(table_widget.rowCount()):
                acr_item = table_widget.item(r, 0)
                defn_item = table_widget.item(r, 1)
                if acr_item and defn_item:
                    acr = acr_item.text().strip()
                    defn = defn_item.text().strip()
                    if acr and defn and (acr not in user_definitions or user_definitions[acr] != defn):
                        user_definitions[acr] = defn
                        made_changes = True

        if made_changes:
            try:
                with open(self.user_acronym_file_path, 'w', encoding='utf-8') as f:
                    for acr, deph in sorted(user_definitions.items()):
                        f.write(f"{acr}\t{deph}\n")
                print(f"User acronym list updated: {self.user_acronym_file_path}")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save user definitions: {e}")

    def create_word_table(self):
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Acronym Table", "", "Word Documents (*.docx)")
        if not file_path:
            return

        try:
            doc = Document()
            doc.add_heading('List of Acronyms and Abbreviations', level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Abbreviation'
            hdr_cells[1].text = 'Definition'

            rows_to_add = []
            for table_widget in [self.likely_table, self.possible_table, self.unlikely_table]:
                for i in range(table_widget.rowCount()):
                    cell_widget = table_widget.cellWidget(i, 2)
                    checkbox = cell_widget.findChild(QCheckBox) if cell_widget else None
                    if checkbox and checkbox.isChecked():
                        acr = table_widget.item(i, 0).text()
                        defn = table_widget.item(i, 1).text() or "---"
                        rows_to_add.append((acr, defn))

            rows_to_add.sort(key=lambda x: x[0].upper())

            for acronym, definition in rows_to_add:
                cells = table.add_row().cells
                cells[0].text = acronym
                cells[1].text = definition

            doc.save(file_path)
            QMessageBox.information(self, "Success", f"Table saved successfully to {file_path}")

        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to generate Word table: {str(e)}")
            print(f"Error generating table: {e}")

    def update_stay_on_top(self):
        parent = self.parent()
        if parent and hasattr(parent, 'stay_on_top_checkbox'):
            flags = self.windowFlags()
            if parent.stay_on_top_checkbox.isChecked():
                self.setWindowFlags(flags | Qt.WindowStaysOnTopHint)
            else:
                self.setWindowFlags(flags & ~Qt.WindowStaysOnTopHint)
            self.show()

    def check_uncheck_all(self, table):
        new_state = None
        for i in range(table.rowCount()):
            cell_widget = table.cellWidget(i, 2)
            checkbox = cell_widget.findChild(QCheckBox) if cell_widget else None
            if checkbox:
                if new_state is None:
                    new_state = not checkbox.isChecked()
                checkbox.setChecked(new_state)

    def add_new_row(self, table):
        """Adds a new, empty row to the specified table."""
        row_position = table.rowCount()
        table.insertRow(row_position)
        table.setItem(row_position, 0, QTableWidgetItem(""))
        table.setItem(row_position, 1, QTableWidgetItem(""))

        checkbox = QCheckBox()
        checkbox.setChecked(True)
        cell_widget = QWidget()
        cell_layout = QHBoxLayout(cell_widget)
        cell_layout.addWidget(checkbox)
        cell_layout.setAlignment(Qt.AlignCenter)
        cell_layout.setContentsMargins(0, 0, 0, 0)
        table.setCellWidget(row_position, 2, cell_widget)

        table.setItem(row_position, 3, QTableWidgetItem(""))
        table.resizeRowToContents(row_position) # Resize this specific new row
        table.editItem(table.item(row_position, 0)) # Start editing

    def add_row_to_current_tab(self):
        """Adds a new row to the table in the currently selected tab."""
        current_widget = self.tab_widget.currentWidget()
        if current_widget:
            table = current_widget.findChild(QTableWidget)
            if table:
                self.add_new_row(table)